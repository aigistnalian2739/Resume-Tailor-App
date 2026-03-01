import os
import re
import tkinter as tk
from tkinter import messagebox, scrolledtext
from docx import Document
from collections import Counter
from datetime import datetime

class ResumeTailor:
    def __init__(self):
        self.stop_words = {'and', 'the', 'with', 'our', 'you', 'will', 'for', 'requirements', 'experience', 'ability', 'work', 'team', 'highly', 'plus', 'preferred'}
        self.industry_keywords = {'python', 'java', 'sql', 'project', 'management', 'agile', 'scrum', 'aws', 'cloud', 'data', 'leadership', 'strategy'}

    def extract_keywords(self, text, num=6):
        words = re.findall(r'\w+', text.lower())
        filtered = [w for w in words if w not in self.stop_words and len(w) > 3]
        counts = Counter(filtered)
        # Give industry-specific terms higher priority
        for word in counts:
            if word in self.industry_keywords:
                counts[word] *= 2
        return ", ".join([w[0].capitalize() for w in counts.most_common(num)])

    def generate(self, company, role, jd_text, years_exp, industry):
        try:
            template = "master_resume.docx"
            if not os.path.exists(template):
                messagebox.showerror("Error", "Place 'master_resume.docx' in this folder first!")
                return

            skills = self.extract_keywords(jd_text)
            doc = Document(template)
            
            replacements = {
                "{{COMPANY}}": company,
                "{{ROLE}}": role,
                "{{SKILLS}}": skills,
                "{{YEARS_EXP}}": years_exp,
                "{{TARGET_INDUSTRY}}": industry,
                "{{DATE}}": datetime.now().strftime("%B %Y")
            }

            # Recursive replacement in paragraphs and tables
            for p in doc.paragraphs:
                for key, val in replacements.items():
                    if key in p.text:
                        p.text = p.text.replace(key, val)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, val in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, val)

            if not os.path.exists("Tailored_Resumes"):
                os.makedirs("Tailored_Resumes")

            output_path = f"Tailored_Resumes/Resume_{company.replace(' ', '_')}.docx"
            doc.save(output_path)
            messagebox.showinfo("Success", f"Resume Generated!\nLocation: {output_path}")
        
        except Exception as e:
            messagebox.showerror("Error", str(e))

def launch_gui():
    root = tk.Tk()
    root.title("AI Resume Tailor")
    root.geometry("600x700")
    
    # Layout UI
    tk.Label(root, text="Target Company:").pack(pady=5)
    ent_comp = tk.Entry(root, width=50); ent_comp.pack()
    
    tk.Label(root, text="Target Role:").pack(pady=5)
    ent_role = tk.Entry(root, width=50); ent_role.pack()

    tk.Label(root, text="Years of Experience / Industry (e.g., 5+ / FinTech):").pack(pady=5)
    ent_meta = tk.Entry(root, width=50); ent_meta.pack()

    tk.Label(root, text="Paste Job Description:").pack(pady=5)
    txt_jd = scrolledtext.ScrolledText(root, height=15, width=65); txt_jd.pack()

    tailor = ResumeTailor()
    
    def on_click():
        meta = ent_meta.get().split("/")
        y = meta[0].strip() if len(meta) > 0 else "X"
        i = meta[1].strip() if len(meta) > 1 else "Tech"
        tailor.generate(ent_comp.get(), ent_role.get(), txt_jd.get("1.0", tk.END), y, i)

    tk.Button(root, text="GENERATE CUSTOM RESUME", bg="#2ecc71", fg="white", 
              font=("Arial", 12, "bold"), command=on_click).pack(pady=20)
    root.mainloop()

if __name__ == "__main__":
    launch_gui()
